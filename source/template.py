from docxtpl import DocxTemplate
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from os.path import exists
from .utils import make_path, make_docx_from_filename
from .pictures import check_and_resize_image
from typing import TYPE_CHECKING
from .models import FileData
from .settings import (
    OUTPUT_FOLDER,
    TEMPLATE_NAME,
    DATE_FORMAT,
    TIME_FORMAT,
    SIGNATURE_FOLDER,
)


if TYPE_CHECKING:
    from models import Person


def create_result_document(filedata: FileData):
    output_filename = make_path(
        OUTPUT_FOLDER, make_docx_from_filename(filedata.filename)
    )
    
    render_template(
        TEMPLATE_NAME,
        {
            "section": getattr(filedata, "section", None),
            "facilityname": getattr(filedata, "facilityname", None),
            "filename": getattr(filedata, "filename", None),
            "documentname": getattr(filedata, "documentname", None),
            "filedate": filedata.modtime.strftime(DATE_FORMAT) if hasattr(filedata, "modtime") else None,
            "filetime": filedata.modtime.strftime(TIME_FORMAT) if hasattr(filedata, "modtime") else None,
            "filehash": getattr(filedata, "filehash", None),
            "documentcode": getattr(filedata, "documentcode", None),
            "documentversion": getattr(filedata, "documentversion", None),
            "filesize": getattr(filedata, "size", None)
        },
        output_filename,
    )

    insert_persons(output_filename, filedata.persons, SIGNATURE_FOLDER)


def render_template(template: str, context: dict[str:str], output: str):
    print(f"Рендеринг файла: {output}")

    doc = DocxTemplate(template)
    doc.render(context)
    doc.save(output)


def insert_persons(
    filename: str, persons: list["Person"], signatures_path: str
) -> None:
    print(f"Добавление персон: {filename}")

    doc = Document(filename)

    table = doc.tables[0].rows[0].cells[0].tables[1]

    for person in persons:
        print(f"\t{person.full_name}")

        new_row = table.rows[-1]
        new_row.cells[0].text = person.job_title
        new_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        new_row.cells[1].text = person.full_name

        picture_path = make_path(signatures_path, person.full_name + ".png")

        if exists(picture_path):
            print(f"\t\tНайдена подпись: {picture_path}")
            check_and_resize_image(picture_path)
            new_row.cells[2].paragraphs[0].add_run().add_picture(picture_path)
            new_row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.add_row()

    doc.save(filename)
